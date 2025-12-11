"""
DocX Jinja Linter Service - Revised Architecture

This module provides comprehensive linting capabilities for Jinja2 templates embedded in .docx files.
It uses docxtpl for consistent extraction and preprocessing, then applies custom parsing logic
for accurate line number detection and unmatched tag analysis.

Workflow:
1. Use docxtpl to extract XML from docx
2. Use docxtpl to process extended docx tags (p, tr, tc, r, etc.)
3. Apply custom logic to lint and find unmatched pair tags
4. Output structured JSON with metadata, formatted input, and detailed errors
5. Optionally convert JSON to markdown and PDF
"""

import re
import time
import json
import tempfile
import os
from typing import List, Dict, Any, Optional, Tuple
from docxtpl import DocxTemplate
from docx import Document
from jinja2 import Environment, TemplateSyntaxError, select_autoescape
from jinja2.exceptions import TemplateError
import logging

from models.schemas import (
    LintResult, LintError, LintWarning, LintSummary, LintOptions,
    LintErrorType, LintWarningType, DocxLinterException, 
    InvalidFileFormatException, TemplateSyntaxException, DocumentExtractionException
)

logger = logging.getLogger(__name__)


class LintResultJson:
    """Structured JSON format for linter results."""
    
    def __init__(self):
        self.metadata = {}
        self.input_data = {}
        self.syntax_errors = []
        self.missing_variables = []
        self.processing_info = {}
    
    def to_dict(self) -> dict:
        return {
            "metadata": self.metadata,
            "input_data": self.input_data,
            "syntax_errors": self.syntax_errors,
            "missing_variables": self.missing_variables,
            "processing_info": self.processing_info
        }


class DocxJinjaLinterService:
    """
    Revised linter service using docxtpl for extraction and custom parsing logic.
    
    Architecture:
    1. Extract XML using docxtpl.get_xml()
    2. Process extended tags using docxtpl.patch_xml()
    3. Parse with custom logic for accurate line numbers
    4. Output structured JSON format
    """
    
    def __init__(self):
        """Initialize the linter service."""
        self.env = Environment(
            autoescape=select_autoescape(['html', 'xml']),
            trim_blocks=True,
            lstrip_blocks=True
        )
        
        # Tags that require matching end tags
        self.paired_tags = {
            'if', 'for', 'with', 'block', 'macro', 'call', 
            'filter', 'trans', 'pluralize', 'raw', 'autoescape'
        }
        
        # Self-contained tags
        self.standalone_tags = {
            'else', 'elif', 'endif', 'endfor', 'endwith', 'endblock', 
            'endmacro', 'endcall', 'endfilter', 'endtrans', 'endpluralize', 
            'endraw', 'endautoescape', 'include', 'import', 'from', 'extends', 
            'break', 'continue', 'set'
        }

    async def lint_docx_file(
        self, 
        file_content: bytes, 
        filename: str,
        options: LintOptions = None
    ) -> LintResult:
        """
        Main linting method using the revised architecture.
        
        Args:
            file_content: Raw bytes of the .docx file
            filename: Original filename for error reporting
            options: Linting configuration options
            
        Returns:
            LintResult containing all errors, warnings, and summary information
        """
        start_time = time.time()
        
        if options is None:
            options = LintOptions()
        
        errors = []
        warnings = []
        
        docxtpl_temp_file_path: Optional[str] = None
        try:
            # Stage 1: Extract plaintext using python-docx (independent of docxtpl)
            logger.info(f"Step 1: Extracting plaintext from {filename} using python-docx")
            structured_text = self._extract_structured_text(file_content, filename)
            
            # Stage 2: Check for unmatched tags in plaintext first
            logger.info(f"Step 2: Checking for unmatched tags in plaintext")
            syntax_errors = self._find_unmatched_tags(structured_text, filename)
            
            # Stage 3: If unmatched tags found, return error report immediately
            if syntax_errors:
                logger.info(f"Step 3: Found {len(syntax_errors)} syntax errors, creating error report")
                errors.extend([self._convert_to_lint_error(err) for err in syntax_errors])
                
                # Create basic input data for error report
                input_data = self._create_basic_input_data(structured_text, filename)
                
                # Save debug output
                await self._save_debug_output_basic(structured_text, filename)
                
                # Skip docxtpl processing and go directly to result creation
                processing_time = (time.time() - start_time) * 1000
                metadata = self._create_metadata(filename, start_time, processing_time)
                
                json_result = self._create_json_result(
                    metadata, input_data, syntax_errors, [], processing_time
                )
                
                # Save intermediate JSON result to file
                await self._save_intermediate_json(json_result, filename)
                
                result = self._convert_json_to_lint_result(
                    json_result, errors, warnings, structured_text, options
                )
                
                # Save intermediate markdown for debugging
                await self._save_intermediate_markdown(result, filename)
                
                logger.info(f"Early termination: {len(errors)} syntax errors found")
                return result
            
            # Stage 4: If syntax is clean, proceed with docxtpl processing
            logger.info(f"Step 4: Syntax clean, proceeding with docxtpl processing")
            doc_template, raw_xml, docxtpl_temp_file_path = self._extract_xml_with_docxtpl(file_content, filename)
            
            # Stage 5: Use docxtpl to process extended docx tags
            logger.info(f"Step 5: Processing extended docx tags with docxtpl")
            processed_xml = self._process_extended_tags(doc_template, raw_xml)
            
            # Stage 6: Create structured input data for JSON
            logger.info(f"Step 6: Creating structured input data")
            input_data = self._create_input_data(raw_xml, processed_xml, structured_text, filename)
            
            # Stage 7: Save debug output for analysis
            logger.info(f"Step 7: Saving debug output for analysis")
            await self._save_debug_output(raw_xml, processed_xml, structured_text, filename)
            
            # Stage 8: Find missing variables using docxtpl
            logger.info(f"Step 8: Finding missing variables")
            if options.check_undefined_vars:
                missing_vars = self._find_missing_variables(doc_template, structured_text)
                for var_info in missing_vars:
                    warnings.append(LintWarning(
                        warning_type=LintWarningType.UNUSED_VARIABLE,
                        message=f"Undefined variable: {var_info['variable']}",
                        line_number=var_info.get('line_number'),
                        suggestion=f"Ensure '{var_info['variable']}' is provided in template data"
                    ))
            
            # Stage 9: Create metadata and processing info
            processing_time = (time.time() - start_time) * 1000
            metadata = self._create_metadata(filename, start_time, processing_time)
            
            # Stage 10: Generate structured JSON result
            json_result = self._create_json_result(
                metadata, input_data, syntax_errors, 
                [self._convert_warning_to_dict(w) for w in warnings], 
                processing_time
            )
            
            # Save intermediate JSON result to file
            await self._save_intermediate_json(json_result, filename)
            
            # Stage 11: Convert to traditional LintResult format
            result = self._convert_json_to_lint_result(
                json_result, errors, warnings, structured_text, options
            )
            
            # Save intermediate markdown for debugging
            await self._save_intermediate_markdown(result, filename)
            
            logger.info(f"Linting completed: {len(errors)} errors, {len(warnings)} warnings")
            return result
            
        except Exception as e:
            logger.error(f"Linting failed for {filename}: {str(e)}")
            return self._create_error_result(e, filename, start_time)
        finally:
            # Keep the docxtpl temp file around for all lazy operations (patch_xml, variable detection, etc.)
            if docxtpl_temp_file_path and os.path.exists(docxtpl_temp_file_path):
                try:
                    os.unlink(docxtpl_temp_file_path)
                except Exception as e:
                    logger.warning(f"Failed to delete temp file {docxtpl_temp_file_path}: {e}")

    def _extract_xml_with_docxtpl(self, file_content: bytes, filename: str) -> Tuple[DocxTemplate, str, str]:
        """
        Step 1: Use docxtpl to extract XML from docx.
        
        Args:
            file_content: Raw bytes of the .docx file
            filename: Original filename for error reporting
            
        Returns:
            Tuple of (DocxTemplate instance, raw XML string, temp file path).
            NOTE: The caller is responsible for deleting the returned temp file path
            after all docxtpl/python-docx lazy operations are complete.
        """
        temp_file_path: Optional[str] = None
        try:
            # Create temporary file - must remain on disk while docxtpl lazily reads it
            with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as temp_file:
                temp_file.write(file_content)
                temp_file_path = temp_file.name

            # Create DocxTemplate instance
            doc_template = DocxTemplate(temp_file_path)
            doc_template.init_docx()

            # Extract raw XML
            raw_xml = doc_template.get_xml()

            logger.debug(f"Successfully extracted XML from {filename}: {len(raw_xml)} characters")
            return doc_template, raw_xml, temp_file_path

        except Exception as e:
            # Clean up the temp file on failure to avoid leaking /tmp files
            if temp_file_path and os.path.exists(temp_file_path):
                try:
                    os.unlink(temp_file_path)
                except Exception as cleanup_err:
                    logger.warning(f"Failed to delete temp file {temp_file_path}: {cleanup_err}")
            raise DocumentExtractionException(
                f"Failed to extract XML from {filename} using docxtpl: {str(e)}"
            )

    def _process_extended_tags(self, doc_template: DocxTemplate, raw_xml: str) -> str:
        """
        Step 2: Use docxtpl to process extended docx tags (p, tr, tc, r, etc.).
        
        Args:
            doc_template: DocxTemplate instance
            raw_xml: Raw XML content
            
        Returns:
            Processed XML with extended tags converted to standard Jinja2
        """
        try:
            # Use docxtpl's patch_xml to process extended tags
            processed_xml = doc_template.patch_xml(raw_xml)
            
            logger.debug(f"Processed extended tags: {len(raw_xml)} -> {len(processed_xml)} characters")
            return processed_xml
            
        except Exception as e:
            logger.warning(f"Failed to process extended tags: {str(e)}")
            # Fallback to raw XML if processing fails
            return raw_xml

    def _extract_structured_text(self, file_content: bytes, filename: str) -> str:
        """
        Step 3: Extract structured text using python-docx to preserve document layout.
        
        Args:
            file_content: Raw bytes of the .docx file
            filename: Document filename for error reporting
            
        Returns:
            Structured text with proper line breaks
        """
        temp_file_path: Optional[str] = None
        try:
            # Create temporary file - must remain on disk while python-docx lazily reads it
            with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as temp_file:
                temp_file.write(file_content)
                temp_file_path = temp_file.name

            doc = Document(temp_file_path)
            full_text = []

            # Extract paragraph text
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():  # Skip empty paragraphs
                    full_text.append(paragraph.text)

            # Extract table text
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        text = cell.text.strip()
                        if text:
                            row_text.append(text)
                    if row_text:
                        full_text.append(' | '.join(row_text))

            structured_text = '\n'.join(full_text)
            logger.debug(f"Extracted structured text: {len(structured_text)} characters, {len(full_text)} lines")
            return structured_text

        except Exception as e:
            logger.error(f"Failed to extract structured text from {filename}: {str(e)}")
            return ""
        finally:
            # Clean up AFTER all operations complete
            if temp_file_path and os.path.exists(temp_file_path):
                try:
                    os.unlink(temp_file_path)
                except Exception as cleanup_err:
                    logger.warning(f"Failed to delete temp file {temp_file_path}: {cleanup_err}")

    def _create_input_data(self, raw_xml: str, processed_xml: str, structured_text: str, filename: str) -> dict:
        """
        Step 4: Create structured input data for JSON output.
        
        Args:
            raw_xml: Original XML content
            processed_xml: Processed XML content
            structured_text: Structured text from python-docx
            filename: Document filename
            
        Returns:
            Structured input data dictionary
        """
        lines = structured_text.split('\n')
        return {
            "filename": filename,
            "raw_xml_size": len(raw_xml),
            "processed_xml_size": len(processed_xml),
            "structured_text_size": len(structured_text),
            "structured_text_lines": len(lines),
            "structured_text_preview": structured_text[:500] + "..." if len(structured_text) > 500 else structured_text,
            "structured_text_full": structured_text,
            "extraction_method": "python-docx + docxtpl",
            "processing_method": "docxtpl.patch_xml + python-docx structure"
        }

    async def _save_debug_output(self, raw_xml: str, processed_xml: str, structured_text: str, filename: str) -> None:
        """
        Step 5: Save debug output files for analysis.
        
        Args:
            raw_xml: Original XML content
            processed_xml: Processed XML content  
            structured_text: Structured text from python-docx
            filename: Document filename
        """
        try:
            base_name = filename.rsplit('.', 1)[0] if '.' in filename else filename
            
            # Save structured text for analysis
            debug_text_path = f"test-data/test-results/debug-{base_name}-structured.txt"
            with open(debug_text_path, 'w', encoding='utf-8') as f:
                f.write("=== STRUCTURED TEXT EXTRACTED WITH python-docx ===\n\n")
                lines = structured_text.split('\n')
                for i, line in enumerate(lines, 1):
                    f.write(f"{i:4d}: {line}\n")
                f.write(f"\n=== SUMMARY ===\n")
                f.write(f"Total lines: {len(lines)}\n")
                f.write(f"Total characters: {len(structured_text)}\n")
            
            # Save processed XML for comparison
            debug_xml_path = f"test-data/test-results/debug-{base_name}-processed.xml"
            with open(debug_xml_path, 'w', encoding='utf-8') as f:
                f.write("=== PROCESSED XML FROM docxtpl ===\n\n")
                f.write(processed_xml)
            
            logger.info(f"Debug files saved: {debug_text_path}, {debug_xml_path}")
            
        except Exception as e:
            logger.error(f"Failed to save debug output: {str(e)}")

    def _create_basic_input_data(self, structured_text: str, filename: str) -> dict:
        """
        Create basic input data for error reports (without docxtpl processing).
        
        Args:
            structured_text: Structured text from python-docx
            filename: Document filename
            
        Returns:
            Basic input data dictionary
        """
        lines = structured_text.split('\n')
        return {
            "filename": filename,
            "structured_text_size": len(structured_text),
            "structured_text_lines": len(lines),
            "structured_text_preview": structured_text[:500] + "..." if len(structured_text) > 500 else structured_text,
            "structured_text_full": structured_text,
            "extraction_method": "python-docx only",
            "processing_method": "plaintext syntax analysis"
        }

    async def _save_debug_output_basic(self, structured_text: str, filename: str) -> None:
        """
        Save basic debug output for syntax error analysis.
        
        Args:
            structured_text: Structured text from python-docx
            filename: Document filename
        """
        try:
            base_name = filename.rsplit('.', 1)[0] if '.' in filename else filename
            
            # Save structured text for analysis
            debug_text_path = f"test-data/test-results/debug-{base_name}-syntax-error.txt"
            with open(debug_text_path, 'w', encoding='utf-8') as f:
                f.write("=== SYNTAX ERROR ANALYSIS - STRUCTURED TEXT ===\n\n")
                lines = structured_text.split('\n')
                for i, line in enumerate(lines, 1):
                    f.write(f"{i:4d}: {line}\n")
                f.write(f"\n=== SUMMARY ===\n")
                f.write(f"Total lines: {len(lines)}\n")
                f.write(f"Total characters: {len(structured_text)}\n")
                f.write(f"Analysis: Syntax errors found in plaintext, docxtpl processing skipped\n")
            
            logger.info(f"Basic debug file saved: {debug_text_path}")
            
        except Exception as e:
            logger.error(f"Failed to save basic debug output: {str(e)}")

    async def _save_intermediate_json(self, json_result: dict, filename: str) -> None:
        """
        Save intermediate JSON result to file for analysis.
        
        Args:
            json_result: The complete JSON result from linting
            filename: Document filename for naming the output file
        """
        try:
            base_name = filename.rsplit('.', 1)[0] if '.' in filename else filename
            
            # Save JSON result
            json_file = f"test-data/test-results/intermediate-{base_name}-linting-result.json"
            
            import json
            with open(json_file, 'w', encoding='utf-8') as f:
                json.dump(json_result, f, indent=2, ensure_ascii=False)
            
            logger.info(f"Intermediate JSON result saved to {json_file}")
            
        except Exception as e:
            logger.error(f"Failed to save intermediate JSON result: {str(e)}")

    async def _save_intermediate_markdown(self, lint_result: 'LintResult', filename: str) -> None:
        """
        Save intermediate markdown for debugging table formatting issues.
        
        Args:
            lint_result: The LintResult object
            filename: Document filename for naming the output file
        """
        try:
            from services.markdown_formatter import create_lint_report_markdown
            
            base_name = filename.rsplit('.', 1)[0] if '.' in filename else filename
            
            # Generate markdown
            markdown_content = create_lint_report_markdown(lint_result, filename)
            
            # Save markdown file
            markdown_file = f"test-data/test-results/debug-{base_name}-report.md"
            with open(markdown_file, 'w', encoding='utf-8') as f:
                f.write(markdown_content)
            
            logger.info(f"Debug markdown saved to {markdown_file}")
            
        except Exception as e:
            logger.error(f"Failed to save debug markdown: {str(e)}")

    def _find_unmatched_tags(self, structured_text: str, filename: str) -> List[dict]:
        """
        Step 6: Apply custom logic to find unmatched pair tags with accurate line numbers.
        
        Args:
            structured_text: Structured text from python-docx
            filename: Document filename for error reporting
            
        Returns:
            List of syntax error dictionaries
        """
        syntax_errors = []
        
        # Split structured text into lines - this preserves document structure
        lines = structured_text.split('\n')
        
        # Note: Removed overly aggressive malformed tag detection
        # docxtpl syntax like {%p if ...%}, {%tr for ...%} etc. is actually VALID
        
        # Join all lines to handle template blocks that span multiple lines
        full_text = '\n'.join(lines)
        
        # Use Jinja2 to parse and find actual syntax errors
        # Replace docxtpl extension tags with regular Jinja2 equivalent
        import re
        preprocessed_text = full_text
        docxtpl_patterns = [
            (r'{%\s*p\s+([^%]*?)%}', r'{% \1 %}'),  # {%p if ...%} -> {% if ...%}
            (r'{%\s*tr\s+([^%]*?)%}', r'{% \1 %}'), # {%tr for ...%} -> {% for ...%}
            (r'{%\s*tc\s+([^%]*?)%}', r'{% \1 %}'), # {%tc if ...%} -> {% if ...%}
            (r'{%\s*r\s+([^%]*?)%}', r'{% \1 %}'),  # {%r if ...%} -> {% if ...%}
        ]
        
        for pattern, replacement in docxtpl_patterns:
            preprocessed_text = re.sub(pattern, replacement, preprocessed_text)
        
        try:
            from jinja2 import Environment, TemplateSyntaxError
            env = Environment()
            
            # Try to parse the preprocessed template
            env.from_string(preprocessed_text)
            
        except TemplateSyntaxError as e:
            # Found a real syntax error - map it back to line numbers
            error_line = e.lineno if e.lineno else 1
            
            # Find the problematic content around the error
            text_lines = full_text.split('\n')
            
            # For "end of template" errors, find the actual unmatched opening tag
            if error_line > len(text_lines) or "Unexpected end of template" in e.message:
                # Look for unmatched opening tags by tracking nesting
                error_content = "End of template - missing closing tag"
                
                # Track all opening and closing tags to find the unmatched one
                stack = []  # Stack to track open tags: [(tag_type, line_number, line_content)]
                
                for i, line in enumerate(text_lines):
                    line_num = i + 1
                    
                    # Check for opening tags (both standard and docxtpl)
                    import re
                    
                    # Find all opening tags in this line
                    opening_patterns = [
                        (r'{%\s*if\s+', 'if'),
                        (r'{%\s*for\s+', 'for'),
                        (r'{%p\s+if\s+', 'if'),  # docxtpl if
                        (r'{%tr\s+for\s+', 'for'),  # docxtpl for
                        (r'{%tc\s+if\s+', 'if'),  # docxtpl table cell if
                        (r'{%r\s+if\s+', 'if'),  # docxtpl row if
                    ]
                    
                    for pattern, tag_type in opening_patterns:
                        if re.search(pattern, line):
                            stack.append((tag_type, line_num, line.strip()))
                    
                    # Check for closing tags
                    if '{% endif %}' in line or '{%p endif %}' in line:
                        if stack and stack[-1][0] == 'if':
                            stack.pop()
                    elif '{% endfor %}' in line or '{%tr endfor %}' in line:
                        if stack and stack[-1][0] == 'for':
                            stack.pop()
                
                # If there are unmatched tags, report the first unmatched one
                if stack:
                    tag_type, line_num, line_content = stack[-1]  # Last unmatched tag
                    error_line = line_num
                    error_content = line_content
            else:
                error_content = text_lines[error_line - 1] if error_line <= len(text_lines) else "Unknown"
            
            syntax_errors.append({
                'type': 'template_syntax_error',
                'line_number': error_line,
                'line_content': error_content.strip(),
                'message': f"Template syntax error: {e.message}",
                'suggestion': "Check for unmatched tags, missing endif/endfor statements, or invalid Jinja2 syntax"
            })
            
            return syntax_errors
            
        except Exception as e:
            # Handle other parsing errors (like unknown tags)
            if 'unknown tag' in str(e).lower():
                # This might be a docxtpl extension tag, try a simpler approach
                pass
            else:
                syntax_errors.append({
                    'type': 'parsing_error',
                    'line_number': 1,
                    'line_content': 'Full template',
                    'message': f"Template parsing error: {str(e)}",
                    'suggestion': "Check template syntax and structure"
                })
                return syntax_errors
        
        # If we get here, no major syntax errors were found by Jinja2
        # The template is syntactically valid, so return no errors
        return syntax_errors

    def _split_xml_into_logical_lines(self, xml_content: str) -> List[str]:
        """
        Split XML content into logical lines for accurate line number reporting.
        
        This attempts to preserve document structure by splitting on:
        1. XML paragraph boundaries (<w:p>)
        2. XML table row boundaries (<w:tr>)
        3. Line break tags (<w:br>)
        4. Actual newlines in the XML
        """
        # Split by meaningful XML boundaries that represent document structure
        logical_lines = []
        
        # First split by paragraph boundaries
        paragraph_parts = re.split(r'(<w:p[^>]*>|</w:p>)', xml_content)
        
        for part in paragraph_parts:
            if not part.strip():
                continue
                
            # Further split by table row boundaries
            row_parts = re.split(r'(<w:tr[^>]*>|</w:tr>)', part)
            
            for row_part in row_parts:
                if not row_part.strip():
                    continue
                    
                # Split by line breaks and actual newlines
                line_parts = re.split(r'(<w:br[^>]*/>|\n)', row_part)
                
                for line_part in line_parts:
                    if line_part.strip():
                        logical_lines.append(line_part)
        
        return logical_lines

    def _extract_jinja_tags(self, line_content: str) -> List[dict]:
        """
        Extract all Jinja2 tags from a line with their types and names.
            
        Returns:
            List of tag dictionaries with type, name, and content
        """
        tags = []
        
        # Pattern for block tags (if, for, etc.)
        block_pattern = r'{%\s*(\w+)(?:\s+([^%]*))?\s*%}'
        
        for match in re.finditer(block_pattern, line_content):
            tag_name = match.group(1)
            tag_args = match.group(2) or ''
            full_content = match.group(0)
            
            if tag_name.startswith('end'):
                tag_type = 'block_end'
            elif tag_name in self.paired_tags:
                tag_type = 'block_start'
            else:
                tag_type = 'standalone'
            
            tags.append({
                'type': tag_type,
                'name': tag_name,
                'args': tag_args.strip(),
                'content': full_content,
                'start_pos': match.start(),
                'end_pos': match.end()
            })
        
        return tags

    def _find_missing_variables(self, doc_template: DocxTemplate, structured_text: str) -> List[dict]:
        """
        Step 7: Find missing variables using docxtpl's built-in method.
        
        Args:
            doc_template: DocxTemplate instance
            structured_text: Structured text content
            
        Returns:
            List of missing variable information
        """
        try:
            # Use docxtpl's method to find undeclared variables
            undeclared_vars = doc_template.get_undeclared_template_variables(self.env)
            
            missing_vars = []
            for var in undeclared_vars:
                # Try to find the line number where this variable is used
                line_number = self._find_variable_line_number(var, structured_text)
                
                missing_vars.append({
                    'variable': var,
                    'line_number': line_number,
                    'type': 'undefined_variable'
                })
            
            return missing_vars
            
        except Exception as e:
            logger.warning(f"Failed to find missing variables: {str(e)}")
            return []

    def _find_variable_line_number(self, variable: str, structured_text: str) -> Optional[int]:
        """
        Find the line number where a variable is first used.
        
        Args:
            variable: Variable name to search for
            structured_text: Structured text content
            
        Returns:
            Line number or None if not found
        """
        lines = structured_text.split('\n')
        
        # Pattern to match variable usage
        var_pattern = r'{{\s*' + re.escape(variable) + r'(?:\s*\|[^}]*)?\s*}}'
        
        for line_num, line_content in enumerate(lines, 1):
            if re.search(var_pattern, line_content):
                return line_num
        
        return None

    def _create_metadata(self, filename: str, start_time: float, processing_time: float) -> dict:
        """Create metadata for the JSON result."""
        return {
            "filename": filename,
            "linter_version": "2.0.0",
            "architecture": "docxtpl-based",
            "timestamp": time.time(),
            "processing_time_ms": round(processing_time, 2)
        }

    def _create_json_result(
        self, 
        metadata: dict, 
        input_data: dict, 
        syntax_errors: List[dict], 
        missing_variables: List[dict],
        processing_time: float
    ) -> dict:
        """Create the structured JSON result."""
        return {
            "metadata": metadata,
            "input_data": input_data,
            "syntax_errors": syntax_errors,
            "missing_variables": missing_variables,
            "processing_info": {
                "total_errors": len(syntax_errors),
                "total_warnings": len(missing_variables),
                "processing_time_ms": round(processing_time, 2),
                "success": len(syntax_errors) == 0
            }
        }

    def _convert_to_lint_error(self, error_dict: dict) -> LintError:
        """Convert dictionary error to LintError object."""
        return LintError(
            line_number=error_dict.get('line_number'),
            error_type=LintErrorType.SYNTAX_ERROR,
            message=error_dict['message'],
            context=error_dict.get('line_content'),
            suggestion=error_dict.get('suggestion')
        )

    def _convert_warning_to_dict(self, warning: LintWarning) -> dict:
        """Convert LintWarning to dictionary."""
        return {
            "type": warning.warning_type.value if warning.warning_type else "unknown",
            "message": warning.message,
            "line_number": warning.line_number,
            "suggestion": warning.suggestion
        }

    def _convert_json_to_lint_result(
        self, 
        json_result: dict, 
        errors: List[LintError], 
        warnings: List[LintWarning],
        structured_text: str,
        options: LintOptions
    ) -> LintResult:
        """Convert JSON result back to traditional LintResult format."""
        
        lines = structured_text.split('\n')
        
        # Create summary
        summary = LintSummary(
            total_errors=len(errors),
            total_warnings=len(warnings),
            template_size=len(structured_text),
            lines_count=len(lines),
            jinja_tags_count=len(re.findall(r'{[%{#].*?[%}#]}', structured_text)),
            completeness_score=100.0 if len(errors) == 0 else max(0.0, 100.0 - (len(errors) * 20)),
            processing_time_ms=json_result['processing_info']['processing_time_ms']
        )
        
        # Create preview
        preview = structured_text[:500] if len(structured_text) > 500 else structured_text
        if len(structured_text) > 500:
            preview += "..."
        
        success = len(errors) == 0 and (not options.fail_on_warnings or len(warnings) == 0)
        
        result = LintResult(
            success=success,
            errors=errors,
            warnings=warnings,
            summary=summary,
            template_content=structured_text if options.verbose else None,
            template_preview=preview
        )
        
        # Store the JSON result for potential later use
        result.json_result = json_result
        
        return result

    def _create_error_result(self, exception: Exception, filename: str, start_time: float) -> LintResult:
        """Create error result when linting fails."""
        processing_time = (time.time() - start_time) * 1000
        
        error = LintError(
            error_type=LintErrorType.DOCUMENT_ERROR,
            message=f"Linting failed: {str(exception)}",
            suggestion="Check document format and try again"
        )
        
        summary = LintSummary(
            total_errors=1,
            total_warnings=0,
            template_size=0,
            lines_count=0,
            jinja_tags_count=0,
            completeness_score=0.0,
            processing_time_ms=round(processing_time, 2)
        )
        
        return LintResult(
            success=False,
            errors=[error],
            warnings=[],
            summary=summary,
            template_content=None,
            template_preview="Error: Could not process document"
        )

    async def export_json_result(self, lint_result: LintResult, output_path: str) -> bool:
        """
        Export the JSON result to a file.
        
        Args:
            lint_result: LintResult containing json_result
            output_path: Path to save JSON file
            
        Returns:
            True if successful, False otherwise
        """
        try:
            if hasattr(lint_result, 'json_result'):
                with open(output_path, 'w') as f:
                    json.dump(lint_result.json_result, f, indent=2)
                return True
            return False
        except Exception as e:
            logger.error(f"Failed to export JSON result: {str(e)}")
            return False